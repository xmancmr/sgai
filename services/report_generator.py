import pandas as pd
from docx import Document
import matplotlib.pyplot as plt
import io

def generate_csv_report(df, path):
    df.to_csv(path, index=False)

def generate_docx_report(df, path, summary=None, plots=None, interpretation=None):
    doc = Document()
    doc.add_heading('Rapport IA SGAI', 0)
    if summary:
        doc.add_heading('Résumé', level=1)
        doc.add_paragraph(summary)
    doc.add_heading('Données', level=1)
    doc.add_paragraph(df.head().to_string())
    if plots:
        for plot_path in plots:
            doc.add_picture(plot_path, width=None)
    if interpretation:
        doc.add_heading('Interprétation', level=1)
        doc.add_paragraph(interpretation)
    doc.save(path)
